from collections.abc import Generator, Iterable, Mapping, Sequence
from contextlib import contextmanager
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from utils.payload import Payload

TITLE_STYLE = "Title"
SUBTITLE_STYLE = "Subtitle"
BULLET_STYLE = "List Bullet"

SECTION_HEADING = 1
ITEM_HEADING = 2


def _get_primary_section(doc: DocumentType):
    """Return the primary section of the document."""
    return doc.sections[0]


def _write_run_into(
    paragraph: Paragraph,
    text: str,
    font_name: str,
    font_size: float,
    **kwargs: dict[str, Any],
) -> Run:
    """Write a run to a paragraph with the specified font name and size.

    Args:
        paragraph (docx.text.paragraph.Paragraph): The paragraph to write the run to.
        text (str): The text for the run.
        font_name (str): The font name for the run.
        font_size (float): The font size for the run.
    """
    run = paragraph.add_run(str(text))
    run.font.name = str(font_name)
    run.font.size = Pt(font_size)

    return run


@contextmanager
def _get_document(path: str | Path) -> Generator[DocumentType, None, None]:
    """Return a Document object from a file path.

    Args:
        path (str | Path): The path to the document file.
    Returns:
        docx.document.Document: The Document object.
    """
    try:
        path = Path(path)
        doc = Document(str(path))

    except (TypeError, ValueError) as e:
        raise ValueError(f"Invalid path provided: {e}") from e

    except FileNotFoundError as e:
        raise FileNotFoundError(f"Document not found at path: {path}") from e

    try:
        yield doc
    finally:
        doc.save(str(path))


def _normalize_margins(
    margins: Sequence[int | float] | Mapping[str, int | float] | None,
) -> dict[str, int | float | None]:
    if margins is None:
        return {}

    keys = ("top", "right", "bottom", "left")

    if isinstance(margins, Mapping):
        return {k: margins.get(k) for k in keys}

    if isinstance(margins, Sequence) and not isinstance(margins, (str, bytes)):
        if len(margins) > 4:
            raise ValueError("Margins sequence cannot have more than 4 values.")
        return {k: margins[i] if i < len(margins) else None for i, k in enumerate(keys)}

    raise TypeError("Margins must be a sequence or mapping type.")


def _set_margins(
    doc: DocumentType,
    margins: Sequence[int | float] | Mapping[str, int | float] | None = None,
) -> None:
    """Set the document margins.

    Args:
        doc (docx.document.Document): The document to set margins for.
        margins (list[int | float] | dict[str, int | float]): The margins to set.
    """
    section = _get_primary_section(doc)

    m = _normalize_margins(margins)

    top = m.get("top") or 1
    right = m.get("right") or top
    bottom = m.get("bottom") or top
    left = m.get("left") or right

    section.top_margin = Inches(top)
    section.right_margin = Inches(right)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)


def _add_name(
    doc: DocumentType,
    text: str,
    font_name: str,
    font_size: float,
    center: bool = True,
) -> None:
    t = doc.add_paragraph(style=TITLE_STYLE)

    if center:
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _write_run_into(t, text, font_name, font_size)


def _add_contact_line(
    doc: DocumentType,
    contacts: list[str | dict[str, str]],
    font_name: str,
    font_size: float,
    center: bool = True,
    sep: str = "|",
) -> None:
    contact_line = doc.add_paragraph(style=SUBTITLE_STYLE)

    if center:
        contact_line.alignment = WD_ALIGN_PARAGRAPH.CENTER

    last_contact_idx = len(contacts) - 1

    for i, contact in enumerate(contacts):
        if isinstance(contact, dict):
            contact_type = contact.get("type").strip()
            contact_value = contact.get("value").strip()

            if not contact_value:
                continue

            if contact_type:
                line = f"{contact_type}: {contact_value}"
            else:
                line = contact_value

        elif isinstance(contact, str):
            line = contact.strip()

            if not line:
                continue

        else:
            continue

        _write_run_into(contact_line, line, font_name, font_size)

        if i < last_contact_idx:
            _write_run_into(contact_line, " " + sep + " ", font_name, font_size)


def _add_summary(
    doc: DocumentType,
    text: str,
    font_name: str,
    font_size: float,
    center: bool = True,
) -> None:
    summary = doc.add_paragraph(style=SUBTITLE_STYLE)

    if center:
        summary.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _write_run_into(summary, text, font_name, font_size)


def _add_section(
    doc: DocumentType,
    section_headering: str,
    items: Iterable[Mapping[str, Any]],
    font_name: str,
    font_size: float,
) -> None:
    doc.add_heading(section_headering, level=SECTION_HEADING)

    def assemble_date(start: str | None = None, end: str | None = None) -> str | None:
        if start and end:
            return f"{start} - {end}"

        return start or end or None

    for item in items:
        title = item.get("title")

        if not title:
            raise ValueError("Section item must have a title.")

        date = assemble_date(item.get("start_date"), item.get("end_date"))

        h = doc.add_heading(level=ITEM_HEADING)
        _write_run_into(h, title, font_name, font_size)

        if date:
            section = _get_primary_section(doc)
            if (
                (width := section.page_width)
                and (left := section.left_margin)
                and (right := section.right_margin)
            ):
                usable_width = width - left - right
                h.paragraph_format.tab_stops.add_tab_stop(
                    usable_width, WD_TAB_ALIGNMENT.RIGHT
                )
                _write_run_into(h, "\t" + date, font_name, font_size)

        if content := item.get("content"):
            p = doc.add_paragraph()
            _write_run_into(p, content, font_name, font_size)

        for bullet in item.get("bullets", []):
            b = doc.add_paragraph(style=BULLET_STYLE)
            _write_run_into(b, bullet, font_name, font_size)


def create_document(doc_path: str | Path) -> None:
    doc = Document()
    doc.save(str(doc_path))


def render(
    doc_path: str | Path,
    payload: Payload,
) -> None:
    """Write the resume document to the specified output path.

    Args:
        doc_path (str | Path): The path to the document to write.
        payload (Payload): The payload containing output path information.
    """
    with _get_document(doc_path) as doc:
        formatting = payload.get("formatting", {})
        content = payload.get("content", {})
        margins = formatting.get("margins", None)
        _set_margins(doc, margins)

        title_text = formatting.get("title_text_style", {})
        _add_name(
            doc,
            content.get("name", "Unnamed"),
            title_text.get("font_name", "Times New Roman"),
            title_text.get("font_size", 16),
            title_text.get("center", True),
        )

        subtitle_text = formatting.get("subtitle_text_style", {})
        _add_contact_line(
            doc,
            content.get("contacts", []),
            subtitle_text.get("font_name", "Times New Roman"),
            subtitle_text.get("font_size", 14),
            subtitle_text.get("center", True),
        )

        summary_text = formatting.get("summary_text_style", {})
        _add_summary(
            doc,
            content.get("summary", ""),
            summary_text.get("font_name", "Times New Roman"),
            summary_text.get("font_size", 11),
            summary_text.get("center", True),
        )

        sections_text = formatting.get("sections_text_style", {})
        sections = content.get("sections", [])

        for section in sections:
            _add_section(
                doc,
                section.get("heading", "Untitled Section"),
                section.get("items", []),
                sections_text.get("font_name", "Times New Roman"),
                sections_text.get("font_size", 11),
            )
